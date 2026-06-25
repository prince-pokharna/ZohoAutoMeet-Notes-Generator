"""
build_docx.py
--------------
Takes the structured notes JSON (from generate_notes.py) for TWO days and
combines them into a single, neatly formatted Word document that looks
like proper hand-made study notes — clear headings, bold key terms,
real bullet lists (not pasted unicode dots), consistent fonts.

Usage:
    python build_docx.py <day1_notes.json> <day2_notes.json> <output.docx>
"""

import sys
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)   # deep blue for headings
BODY_FONT = "Calibri"
HEADING_FONT = "Calibri"


def set_run_font(run, font_name=BODY_FONT, size=11, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # Ensure East-Asian font fallback doesn't break formatting on some systems
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_bullet_paragraph(doc, text, bold_terms=None):
    """Adds a real bulleted list item using Word's built-in List Bullet style
    (never paste a unicode • character manually)."""
    p = doc.add_paragraph(style="List Bullet")
    add_rich_text(p, text)
    return p


def add_rich_text(paragraph, text):
    """
    Parses **bold** markdown-style markers in a line of text and adds
    properly bolded runs to the paragraph, instead of a flat text blob.
    """
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_day_section(doc, notes: dict, day_label: str):
    # --- Day banner heading ---
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = banner.add_run(f"{day_label}  —  {notes.get('title', 'Session Notes')}")
    set_run_font(run, font_name=HEADING_FONT, size=18, bold=True, color=ACCENT_COLOR)

    # Underline rule beneath the banner for a clean "notebook divider" look
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_after = Pt(12)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Sections ---
    for section in notes.get("sections", []):
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(14)
        heading_p.paragraph_format.space_after = Pt(6)
        run = heading_p.add_run(section.get("heading", ""))
        set_run_font(run, font_name=HEADING_FONT, size=14, bold=True, color=ACCENT_COLOR)

        content = section.get("content", "")
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("• "):
                add_bullet_paragraph(doc, line[2:].strip())
            else:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(8)
                add_rich_text(para, line)

    # --- Key takeaways box ---
    takeaways = notes.get("key_takeaways", [])
    if takeaways:
        kt_heading = doc.add_paragraph()
        kt_heading.paragraph_format.space_before = Pt(16)
        run = kt_heading.add_run("Key Takeaways")
        set_run_font(run, font_name=HEADING_FONT, size=13, bold=True, color=RGBColor(0x00, 0x66, 0x33))

        for point in takeaways:
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, point)


def build_combined_docx(day1_notes_path: str, day2_notes_path: str, output_path: str,
                         day1_label: str = "Day 1", day2_label: str = "Day 2"):
    doc = Document()

    # Base document-wide font setup
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    # Page setup with comfortable margins (looks more like notes, less like a memo)
    section = doc.sections[0]
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    # --- Cover title ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Session Notes")
    set_run_font(run, font_name=HEADING_FONT, size=24, bold=True, color=ACCENT_COLOR)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(f"{day1_label} & {day2_label}")
    set_run_font(run, size=13, bold=False, color=RGBColor(0x44, 0x44, 0x44))
    subtitle_p.paragraph_format.space_after = Pt(24)

    # --- Day 1 ---
    day1_notes = json.loads(Path(day1_notes_path).read_text(encoding="utf-8"))
    add_day_section(doc, day1_notes, day1_label)

    # Page break between the two days
    doc.add_page_break()

    # --- Day 2 ---
    day2_notes = json.loads(Path(day2_notes_path).read_text(encoding="utf-8"))
    add_day_section(doc, day2_notes, day2_label)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Combined notes saved to: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 4:
        print("Usage: python build_docx.py <day1_notes.json> <day2_notes.json> <output.docx>")
        sys.exit(1)

    build_combined_docx(sys.argv[1], sys.argv[2], sys.argv[3])
